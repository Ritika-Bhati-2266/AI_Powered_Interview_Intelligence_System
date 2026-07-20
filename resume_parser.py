"""
Resume Parser Module
Extracts text content from PDF and DOCX resume files.
Uses PyPDF2 for PDFs and python-docx for DOCX files.
Provides basic skill extraction from text content.
"""

import re
import os


def extract_text_from_pdf(filepath: str) -> str:
    """
    Extract text content from a PDF file.
    
    Args:
        filepath: Absolute or relative path to PDF file
    
    Returns:
        Extracted text as a single string
    
    Raises:
        FileNotFoundError: If the file doesn't exist
        ValueError: If the file is not a valid PDF
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Resume file not found: {filepath}")

    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 is required. Install with: pip install PyPDF2")

    try:
        reader = PdfReader(filepath)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        full_text = "\n".join(text_parts)

        if not full_text.strip():
            # Some PDFs have extractable text issues
            raise ValueError("PDF appears to be a scanned image or has no extractable text")

        return full_text.strip()

    except Exception as e:
        if isinstance(e, (FileNotFoundError, ValueError, ImportError)):
            raise
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(filepath: str) -> str:
    """
    Extract text content from a DOCX file.
    
    Args:
        filepath: Absolute or relative path to DOCX file
    
    Returns:
        Extracted text as a single string
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"Resume file not found: {filepath}")

    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required. Install with: pip install python-docx")

    try:
        doc = Document(filepath)
        text_parts = []

        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts).strip()

    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_text(filepath: str) -> str:
    """
    Extract text from a resume file (auto-detect format by extension).
    
    Args:
        filepath: Path to resume (PDF or DOCX)
    
    Returns:
        Extracted plain text
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == '.pdf':
        return extract_text_from_pdf(filepath)
    elif ext == '.docx':
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload a PDF or DOCX file.")


def extract_skills(text: str) -> list:
    """
    Extract technical skills from resume text using keyword matching.
    Detects programming languages, frameworks, tools, and technologies.
    
    Args:
        text: Plain text extracted from resume
    
    Returns:
        Sorted list of identified skills (lowercase, deduplicated)
    """
    if not text:
        return []

    text_lower = text.lower()

    # Comprehensive skill keyword list organized by category
    skill_keywords = {
        # Programming Languages
        "python", "javascript", "typescript", "java", "c++", "c#", "ruby", "go", "golang",
        "rust", "swift", "kotlin", "scala", "php", "perl", "r", "matlab", "dart", "lua",
        "bash", "shell", "sql", "graphql", "html", "css", "sass", "less", "julia",
        "c", "objective-c", "elixir", "haskell", "clojure", "solidity",

        # Frontend Frameworks
        "react", "angular", "vue", "vue.js", "svelte", "next.js", "nuxt", "gatsby",
        "jquery", "bootstrap", "tailwind", "material-ui", "chakra-ui", "redux",
        "webpack", "vite", "parcel", "electron", "react native", "flutter",

        # Backend Frameworks
        "django", "flask", "fastapi", "spring", "spring boot", "ruby on rails",
        "express", "node.js", "laravel", "asp.net", "gin", "echo", "fiber",
        "ktor", "actix", "rocket",

        # Databases
        "postgresql", "postgres", "mysql", "mongodb", "redis", "sqlite",
        "oracle", "sql server", "mariadb", "cassandra", "dynamodb", "firebase",
        "elasticsearch", "neo4j", "couchdb", "influxdb", "supabase",

        # Cloud & DevOps
        "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s",
        "jenkins", "github actions", "gitlab ci", "circleci", "terraform",
        "ansible", "puppet", "chef", "helm", "argocd", "prometheus", "grafana",
        "nginx", "apache", "cloudflare", "heroku", "vercel", "netlify",

        # Data Science & ML
        "tensorflow", "pytorch", "keras", "scikit-learn", "pandas", "numpy",
        "jupyter", "spark", "hadoop", "airflow", "mlflow", "dvc",
        "opencv", "nltk", "hugging face", "langchain", "llama", "ollama",
        "machine learning", "deep learning", "nlp", "computer vision",
        "data science", "data analysis", "statistics",

        # Tools & Platforms
        "git", "github", "gitlab", "bitbucket", "jira", "confluence",
        "figma", "sketch", "photoshop", "trello", "asana", "notion",
        "slack", "discord", "postman", "insomnia", "swagger",
        "linux", "unix", "macos", "windows", "vim", "vscode", "intellij",
        "yarn", "npm", "pnpm", "pip", "maven", "gradle",

        # Testing
        "jest", "mocha", "chai", "pytest", "cypress", "selenium",
        "playwright", "junit", "unittest", "rspec", "testng",

        # Concepts
        "rest api", "restful", "microservices", "api", "graphql",
        "ci/cd", "agile", "scrum", "tdd", "devops", "serverless",
        "event-driven", "message queue", "rabbitmq", "kafka",
        "unit testing", "integration testing", "e2e testing",
    }

    found_skills = set()
    skill_pattern_lower = {s.lower() for s in skill_keywords}

    # Check each skill keyword
    for skill in skill_pattern_lower:
        # Use word boundary matching to avoid false positives
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text_lower):
            found_skills.add(skill)

    # Also extract capitalized proper nouns that appear with high frequency
    # (company names, specific technologies not in our list)
    words = text_lower.split()

    return sorted(found_skills)


def extract_name(text: str) -> str:
    """
    Attempt to extract candidate name from resume text.
    Looks for patterns like "Name" at the top of the document.
    
    Args:
        text: Plain text from resume
    
    Returns:
        Extracted name or empty string
    """
    if not text:
        return ""

    lines = text.strip().split('\n')
    
    # Usually the first non-empty line is the name
    for line in lines[:10]:
        line = line.strip()
        if line and len(line) < 40 and not re.match(r'^[\d\s@#$%^&*()]+$', line):
            # Skip obvious non-name lines
            if not any(word in line.lower() for word in
                       ['resume', 'cv', 'curriculum', 'vitae', 'email', '@', 'phone',
                        'linkedin', 'github', 'address', 'summary', 'profile']):
                # Check it looks like a name (2-4 words, mostly alpha)
                words = line.split()
                if 1 < len(words) <= 4 and all(w.isalpha() for w in words if w):
                    return line.strip()

    return ""


def extract_email(text: str) -> str:
    """Extract email address from text."""
    match = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', text)
    return match.group(0) if match else ""


def parse_resume(filepath: str) -> dict:
    """
    Complete resume parsing: extract text, detect format, identify skills,
    and return structured information.
    
    Args:
        filepath: Path to resume file
    
    Returns:
        Dict with keys: text, skills, name, email, format, success, error
    """
    result = {
        "text": "",
        "skills": [],
        "name": "",
        "email": "",
        "format": "",
        "success": False,
        "error": None,
    }

    try:
        # Detect format
        ext = os.path.splitext(filepath)[1].lower()
        if ext == '.pdf':
            result["format"] = "pdf"
            result["text"] = extract_text_from_pdf(filepath)
        elif ext == '.docx':
            result["format"] = "docx"
            result["text"] = extract_text_from_docx(filepath)
        else:
            result["error"] = f"Unsupported format: {ext}"
            return result

        # Extract information
        result["skills"] = extract_skills(result["text"])
        result["name"] = extract_name(result["text"])
        result["email"] = extract_email(result["text"])
        result["success"] = True

    except FileNotFoundError as e:
        result["error"] = str(e)
    except ValueError as e:
        result["error"] = str(e)
    except ImportError as e:
        result["error"] = str(e)
    except Exception as e:
        result["error"] = f"Unexpected error parsing resume: {str(e)}"

    return result
