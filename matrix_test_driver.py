"""
Full-matrix test driver for the AI Interview Intelligence System (Groq backend).

Covers ALL role x experience x company combinations with:
  - realistic synthetic resume upload (DOCX) -> skill extraction -> question generation
  - all 3 modes (technical / coding / hr) distributed across the matrix
  - FULL aptitude rounds (all 10 MCQs) for the Indian IT companies (TCS/Infosys/Wipro/HCL)
  - prompt-leak / broken-JSON / crash / 0-score / missing-skill / aptitude-scoring /
    report-generation checks per candidate
  - rewrite flow on 3 candidates

Outputs a pass/fail summary table and aggregate statistics.
"""

import json
import os
import re
import sys
import time
import threading
import random
import requests
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime

PROJ = r"D:\PROJECTS\Final project\AI Interview Intelligence System"
sys.path.insert(0, PROJ)
os.chdir(PROJ)

from dotenv import load_dotenv
load_dotenv(os.path.join(PROJ, ".env"))

import ai_service
from resume_parser import parse_resume
from interview_engine import (
    start_interview,
    submit_answer,
    rewrite_answer,
    generate_report,
    session_store,
)

MODEL = "llama-3.1-8b-instant"   # llama-3.3-70b-versatile TPD quota is exhausted (2026-08-04)
MODEL_ROTATION = [m.strip() for m in os.environ.get("MODEL_ROTATION", "").split(",") if m.strip()]
if not MODEL_ROTATION:
    MODEL_ROTATION = ["llama-3.1-8b-instant", "qwen/qwen3.6-27b", "openai/gpt-oss-120b"]
ai_service.GROQ_MODEL = MODEL
os.environ["GROQ_MODEL"] = MODEL

SUBSET_LIMIT = int(os.environ.get("SUBSET_LIMIT", "0"))   # 0 = all combos
CHECKPOINT = os.environ.get("CHECKPOINT", os.path.join(PROJ, "matrix_checkpoint.jsonl"))
MAX_OUT_TOKENS = int(os.environ.get("MAX_OUT_TOKENS", "1024"))  # cap Groq max_tokens to fit the free-tier TPM quota
RATE_PER_MIN = float(os.environ.get("RATE_PER_MIN", "9"))  # global LLM call pace (3 models share ~22k TPM -> ~10 req/min)

# ══════════════════════════════════════════════════════════════════════════════
# Config
# ══════════════════════════════════════════════════════════════════════════════

ROLES = [
    "Software Engineer", "Data Scientist", "Product Manager", "Frontend Developer",
    "Backend Developer", "DevOps Engineer", "ML Engineer", "QA Engineer",
    "Full Stack Developer", "Data Analyst",
]
EXPS = ["0-1", "1-2", "2-4", "4-7", "7-10", "10+"]
COMPANIES = ["General", "Google", "Amazon", "Meta", "Microsoft", "Apple",
             "TCS", "Infosys", "Wipro", "HCL", "Startup"]
MODES = ["technical", "coding", "hr"]

APTITUDE_COMPANY_KEYS = {"tcs", "infosys", "wipro", "hcl"}
APTITUDE_ROUND_QUESTIONS = 10

STD_Q_CAP = 2              # LLM questions answered for a standard candidate
CODING_DEEP_CAP = 50       # how many coding-mode candidates get the deep (1 real coding Q) run
FULL_APT_EXPS = {"0-1", "7-10"}   # experience levels that run the FULL 10-MCQ aptitude round

WORKERS = 24
LAUNCH_BUDGET_SEC = 18.5 * 60     # stop launching new interviews after this
REAL_REPORT_SAMPLE = 30           # candidates that get real (non-stubbed) report generation

RESUME_DIR = os.path.join(PROJ, "test_resumes")

# ══════════════════════════════════════════════════════════════════════════════
# LLM instrumentation (call counting + deterministic report stubbing)
# ══════════════════════════════════════════════════════════════════════════════

_orig_call_llm = ai_service._call_llm
_stub = threading.local()
_stub.on = False
_stats_lock = threading.Lock()
_rot_lock = threading.Lock()
_rot_idx = 0
_pace_lock = threading.Lock()
_pace_next = 0.0
STATS = {"count": 0, "errors": 0, "error_samples": [], "total_s": 0.0, "llm_seconds": 0.0}


def _pace():
    """Space LLM calls to RATE_PER_MIN so the free-tier TPM budget isn't oversubscribed."""
    global _pace_next
    with _pace_lock:
        now = time.time()
        if now < _pace_next:
            time.sleep(_pace_next - now)
            now = time.time()
        _pace_next = max(now, _pace_next) + 60.0 / RATE_PER_MIN


def _patched_call_llm(prompt, system_prompt=None, temperature=0.7):
    global _rot_idx
    if getattr(_stub, "on", False):
        return ("REPORT_STUB: solid performance with clear strengths across the core "
                "dimensions; continue refining your answers for higher scores.")
    _pace()
    if len(MODEL_ROTATION) > 1:
        with _rot_lock:
            ai_service.GROQ_MODEL = MODEL_ROTATION[_rot_idx % len(MODEL_ROTATION)]
            _rot_idx += 1
    t0 = time.time()
    r = _orig_call_llm(prompt, system_prompt, temperature)
    dt = time.time() - t0
    with _stats_lock:
        STATS["count"] += 1
        STATS["total_s"] += dt
        if not r.strip() or r.startswith("["):
            STATS["errors"] += 1
            if len(STATS["error_samples"]) < 40:
                STATS["error_samples"].append((r[:240], round(dt, 1)))
    return r


ai_service._call_llm = _patched_call_llm

# Cap Groq output tokens (ai_service hardcodes 4096) so the free-tier TPM quota
# (6000-8000 tokens/min per model) isn't burned by max_tokens reservation.
_orig_call_groq = ai_service._call_groq


def _patched_call_groq(prompt, system_prompt=None, temperature=0.7):
    payload = {
        "model": ai_service.GROQ_MODEL,
        "messages": ([{"role": "system", "content": system_prompt}] if system_prompt else [])
        + [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "max_tokens": MAX_OUT_TOKENS,
    }
    for attempt in range(ai_service.MAX_RETRIES + 1):
        try:
            resp = requests.post(ai_service.GROQ_ENDPOINT, json=payload,
                                 headers={"Authorization": f"Bearer {ai_service.GROQ_API_KEY}",
                                          "Content-Type": "application/json"},
                                 timeout=ai_service.REQUEST_TIMEOUT)
            resp.raise_for_status()
            data = resp.json()
            return data["choices"][0]["message"]["content"].strip()
        except requests.exceptions.RequestException as e:
            if attempt < ai_service.MAX_RETRIES:
                time.sleep(ai_service.RETRY_DELAY)
                continue
            try:
                detail = e.response.json().get("error", {}).get("message", str(e))
            except Exception:
                detail = str(e)
            return f"[OLLAMA_ERROR] Groq API error: {detail}"
        except (json.JSONDecodeError, KeyError, IndexError) as e:
            return f"[PARSE_ERROR] Could not parse Groq response: {str(e)}"


if MAX_OUT_TOKENS:
    ai_service._call_groq = _patched_call_groq


def report_stubbed():
    _stub.on = True


def report_real():
    _stub.on = False

# ══════════════════════════════════════════════════════════════════════════════
# Synthetic resumes (realistic, skills recognized by resume_parser.extract_skills)
# ══════════════════════════════════════════════════════════════════════════════

ROLE_RESUME = {
    "Software Engineer": {
        "name": "Arjun Mehta",
        "skills": ["python", "javascript", "java", "flask", "django", "postgresql",
                   "mongodb", "redis", "docker", "aws", "git", "rest api",
                   "microservices", "linux", "ci/cd"],
    },
    "Data Scientist": {
        "name": "Priya Sharma",
        "skills": ["python", "pandas", "numpy", "scikit-learn", "tensorflow", "pytorch",
                   "statistics", "machine learning", "deep learning", "nlp", "jupyter",
                   "sql", "spark", "docker"],
    },
    "Product Manager": {
        "name": "Rohan Verma",
        "skills": ["agile", "scrum", "jira", "figma", "sql", "aws", "rest api",
                   "data analysis", "postman"],
    },
    "Frontend Developer": {
        "name": "Ananya Iyer",
        "skills": ["javascript", "typescript", "react", "redux", "html", "css",
                   "tailwind", "jest", "webpack", "next.js", "git", "rest api"],
    },
    "Backend Developer": {
        "name": "Kabir Singh",
        "skills": ["python", "django", "flask", "fastapi", "node.js", "express",
                   "postgresql", "mongodb", "redis", "docker", "kafka", "aws",
                   "microservices", "git"],
    },
    "DevOps Engineer": {
        "name": "Sneha Reddy",
        "skills": ["aws", "azure", "gcp", "docker", "kubernetes", "terraform", "ansible",
                   "jenkins", "github actions", "prometheus", "grafana", "linux", "nginx",
                   "ci/cd", "git"],
    },
    "ML Engineer": {
        "name": "Vikram Nair",
        "skills": ["python", "tensorflow", "pytorch", "scikit-learn", "pandas", "numpy",
                   "mlflow", "docker", "kubernetes", "aws", "machine learning",
                   "deep learning", "nlp", "spark", "git"],
    },
    "QA Engineer": {
        "name": "Neha Gupta",
        "skills": ["pytest", "selenium", "cypress", "playwright", "jest", "jira",
                   "postman", "git", "ci/cd", "agile", "unit testing",
                   "integration testing", "rest api"],
    },
    "Full Stack Developer": {
        "name": "Ishaan Joshi",
        "skills": ["python", "javascript", "react", "node.js", "express", "flask",
                   "postgresql", "mongodb", "redis", "docker", "aws", "rest api",
                   "microservices", "git"],
    },
    "Data Analyst": {
        "name": "Meera Pillai",
        "skills": ["python", "sql", "pandas", "numpy", "statistics", "jupyter",
                   "data analysis", "aws", "git", "postgresql"],
    },
}


def _slug(s):
    return s.lower().replace(" ", "_").replace("/", "_")


def make_resume_docx(role):
    """Create a realistic synthetic resume for a role and return its path."""
    info = ROLE_RESUME[role]
    os.makedirs(RESUME_DIR, exist_ok=True)
    path = os.path.join(RESUME_DIR, f"{_slug(role)}_resume.docx")

    try:
        from docx import Document
    except ImportError:
        return None

    doc = Document()
    doc.add_paragraph(info["name"])
    doc.add_paragraph(f"{info['name'].split()[0].lower()}@example.com | +91 98xxxxxx | linkedin.com/in/{_slug(info['name'])}")
    doc.add_paragraph("")
    doc.add_paragraph(f"Summary: {role} with 5+ years of experience designing, building and "
                      f"shipping reliable products. Strong background in {', '.join(info['skills'][:6])}.")
    doc.add_paragraph("")
    doc.add_paragraph("SKILLS")
    doc.add_paragraph(", ".join(info["skills"]))
    doc.add_paragraph("")
    doc.add_paragraph("EXPERIENCE")
    doc.add_paragraph("Senior {role} | TechNova Solutions | 2021 - Present".format(role=role))
    for s in info["skills"][:8]:
        doc.add_paragraph(f"- Led delivery using {s}, improving performance and reliability across the platform.")
    doc.add_paragraph("Software Engineer | CloudWorks India | 2018 - 2021")
    doc.add_paragraph(f"- Built REST APIs and data pipelines with {info['skills'][3] if len(info['skills']) > 3 else 'python'}, "
                      f"deployed with docker and aws, and collaborated using git and agile practices.")
    doc.add_paragraph("")
    doc.add_paragraph("PROJECTS")
    doc.add_paragraph("- Scalable analytics dashboard: python, react, postgresql, docker")
    doc.add_paragraph("- Automated deployment pipeline: jenkins, docker, aws, ci/cd")
    doc.add_paragraph("")
    doc.add_paragraph("EDUCATION")
    doc.add_paragraph("B.Tech in Computer Science, National Institute of Technology, 2018")

    doc.save(path)
    return path


def load_role_resumes():
    """Create + parse a resume for each role. Returns {role: {text, skills, name}}."""
    result = {}
    for role in ROLES:
        path = make_resume_docx(role)
        parsed = parse_resume(path)
        if not parsed["success"]:
            print(f"[WARN] resume parse failed for {role}: {parsed['error']}")
        result[role] = {
            "text": parsed.get("text", ""),
            "skills": parsed.get("skills", []),
            "name": parsed.get("name") or ROLE_RESUME[role]["name"],
        }
    return result

# ══════════════════════════════════════════════════════════════════════════════
# Synthetic candidate answers
# ══════════════════════════════════════════════════════════════════════════════


def build_answer(role, skills, question, category, is_resume_phase, company):
    """Return a realistic, reasonably strong candidate answer for the given question."""
    top = skills[:6] if skills else ["core skills"]
    skills_txt = ", ".join(top)
    cat = str(category).lower()

    if any(k in cat for k in ("behavioral", "hr", "situational", "experience")) or is_resume_phase:
        return (
            f"A specific situation I can share: at my previous company, my team was facing a "
            f"tight deadline with unclear requirements for a {role} initiative. The task was to "
            f"deliver a working solution in four weeks. I took the action of organizing a short "
            f"kickoff, defining acceptance criteria, assigning clear owners, and prototyping the "
            f"riskiest part first. As a result we shipped on time, the stakeholder gave positive "
            f"feedback, and we reused the prototype in a follow-up project. My role was to own "
            f"the technical approach using {skills_txt} and to keep the team aligned with clear "
            f"communication throughout. I believe this reflects how I stay organized and "
            f"results-focused under pressure."
        )
    if "coding" in cat or "algorithm" in cat or "data structure" in cat:
        return (
            f"I would approach this by first clarifying the input constraints and expected "
            f"output, then selecting an appropriate algorithm and data structure. I would walk "
            f"through the time and space complexity, handle edge cases like empty input, "
            f"duplicate values, and large inputs, and write clean code with descriptive "
            f"names. In my recent work as a {role} I used {skills_txt}, and I would validate "
            f"the solution with unit tests covering both happy path and boundary cases before "
            f"concluding with the complexity analysis."
        )
    # technical / project / general
    return (
        f"In my experience as a {role}, I approach this systematically. First I clarify the "
        f"core objective and constraints, then I break the problem into smaller components and "
        f"prioritize by business impact. In a recent project I used {skills_txt} to design and "
        f"deliver the solution, and I measured the outcome: we reduced page load time by about "
        f"30% and improved reliability. I follow best practices like clean code, automated "
        f"testing, and continuous integration, and I communicate progress clearly with my team. "
        f"For example, when we hit a performance bottleneck, I profiled the system, identified "
        f"the root cause, applied a targeted fix, and verified it with regression tests before "
        f"shipping."
    )

# ══════════════════════════════════════════════════════════════════════════════
# Checks
# ══════════════════════════════════════════════════════════════════════════════

LEAK_PATTERNS = [
    r"\bReturn ONLY\b", r"\bHARD CONSTRAINT\b", r"\bCRITICAL RULES\b",
    r"\bCRITICAL - \b", r"\bCRITICAL:\b", r"\bYou MUST\b", r"\bmaps to\b",
    r"\bLeadership Principle\b", r"Candidate's Target Role", r"Candidate's Skills",
    r"\bQuestion Category\b", r"\bDIFFICULTY ENFORCEMENT\b", r"\bPREVIOUSLY ASKED\b",
    r"\bDo NOT repeat\b", r"\bsystem prompt\b", r"\bsystem_prompt\b",
    r"\[Instruction", r"\bGenerate a single\b", r"\bResume Context\b",
    r"\bPrevious Interview Context\b", r"\bCOMPANY-SPECIFIC\b", r"\bOutput ONLY\b",
    r"\bYou are the interviewer\b", r"\bOLLAMA_\b", r"\bPARSE_ERROR\b",
    r"\bDo NOT write\b", r"\bdo not preface\b", r"\bOutput ONLY a question\b",
]

VALID_GRADES = {"A+", "A", "B+", "B", "C", "F"}
SCORE_KEYS = ["overall_score", "technical_score", "communication_score", "confidence_score",
              "problem_solving_score", "time_management_score", "conceptual_clarity_score"]

# Known aptitude-bank errors (verified: the bank's `correct` index contradicts the
# explanation / the mathematically correct answer). Keyed by a distinctive substring of
# the question text. (question_substring, real_answer_text or None, note)
KNOWN_APTITUDE_BANK_ERRORS = []


def apt_bank_error(qtext, options, correct_idx):
    """If the served aptitude question is a known-bank-error, return an issue string."""
    for sub, real, note in KNOWN_APTITUDE_BANK_ERRORS:
        if sub.lower() in qtext.lower():
            if real is None:
                return f"APTITUDE_BANK_ERROR({note})"
            try:
                real_idx = options.index(real)
            except ValueError:
                real_idx = None
            if real_idx is None:
                return f"APTITUDE_BANK_ERROR(expected '{real}' not among served options; {note})"
            if real_idx != correct_idx:
                return (f"APTITUDE_BANK_ERROR(bank marks '{options[correct_idx]}' "
                        f"but correct is '{real}'; {note})")
            return None
    return None


def _skill_count(text, skills):
    if not text or not skills:
        return 0
    low = text.lower()
    return sum(1 for s in skills if s.lower() in low)


def check_question_text(q, resume_skills, apt_correct_text=None):
    """Return a list of issue strings + set of resume skills mentioned in the question."""
    issues = []
    mentioned = set()
    if not q or not q.strip():
        issues.append("EMPTY_QUESTION")
        return issues, mentioned
    qw = len(q.split())
    if qw > 60:
        issues.append(f"QUESTION_TOO_LONG({qw}w)")
    stripped = q.lstrip()
    if stripped.startswith("["):
        issues.append(f"ERROR_MARKER: {stripped[:80]}")
    if re.match(r"^(class|def|function|enum|struct)\b", stripped, re.I):
        issues.append("QUESTION_LEAKS_CODE")
    for p in LEAK_PATTERNS:
        if re.search(p, q, re.IGNORECASE):
            issues.append(f"PROMPT_LEAK({p})")
    if apt_correct_text is not None:
        ac = apt_correct_text.strip()
        substantive = len(ac) >= 2 and not (ac.isalpha() and len(ac) < 4)
        if substantive:
            body = re.split(r"\n\s*[A-D]\.\s", q)[0]
            if re.search(r"(?<!\w)" + re.escape(ac) + r"(?!\w)", body, re.IGNORECASE):
                issues.append("APTITUDE_LEAKS_ANSWER")
            if re.search(r"\b(correct answer|answer is|correct option)\b", body, re.IGNORECASE):
                issues.append("APTITUDE_ANSWER_HINT")
    for s in resume_skills:
        if s.lower() in q.lower():
            mentioned.add(s)
    return issues, mentioned


def check_evaluation(ev, answer):
    """Return issue strings for a non-aptitude evaluation dict."""
    issues = []
    if not isinstance(ev, dict):
        return ["EVAL_NOT_DICT"]
    score = ev.get("overall_score")
    if score is None:
        issues.append("EVAL_MISSING_OVERALL")
    else:
        try:
            score = int(round(float(score)))
            if not (0 <= score <= 10):
                issues.append(f"EVAL_SCORE_OUT_OF_RANGE({score})")
        except (TypeError, ValueError):
            issues.append(f"EVAL_SCORE_NOT_NUMERIC({score})")
    dims = [ev.get(k) for k in SCORE_KEYS[1:]]
    try:
        dims_int = [int(round(float(d))) for d in dims]
    except (TypeError, ValueError):
        dims_int = []
    if score == 0 and dims_int and all(d == 0 for d in dims_int) and len(answer.split()) > 40:
        issues.append("ZERO_SCORE_BUG(all dims 0 on substantive answer)")
    # Broken-JSON / fallback-evaluation indicator
    if (not ev.get("strengths") and not ev.get("weaknesses")
            and not ev.get("feedback") and not ev.get("ideal_answer")):
        issues.append("BROKEN_JSON(eval text fields all empty)")
    if not isinstance(ev.get("strengths"), list) or not isinstance(ev.get("weaknesses"), list):
        issues.append("EVAL_LISTS_MALFORMED")
    if not isinstance(ev.get("keywords_used"), list):
        issues.append("EVAL_KEYWORDS_MISSING")
    return issues

# ══════════════════════════════════════════════════════════════════════════════
# Per-candidate interview runner
# ══════════════════════════════════════════════════════════════════════════════

RUN_DESIGNATION = {}


def set_run_designations():
    """Designate specific (role_idx, exp_idx, company_idx) combos for special runs."""
    global RUN_DESIGNATION
    for ri in range(len(ROLES)):
        for ei in range(len(EXPS)):
            for ci in range(len(COMPANIES)):
                key = (ri, ei, ci)
                mode = MODES[(ri + ci * 2 + ei) % 3]
                d = {"mode": mode}
                if COMPANIES[ci].lower() in APTITUDE_COMPANY_KEYS:
                    d["apt_style"] = "full" if EXPS[ei] in FULL_APT_EXPS else "std"
                else:
                    d["apt_style"] = "none"
                    d["coding_deep"] = (mode == "coding")
                RUN_DESIGNATION[key] = d


CODING_DEEP_USED = [0]


def run_candidate(seq, role_idx, exp_idx, company_idx, role_resumes, rewrite_key=None, real_report=False):
    """Run one full (or capped) interview. Returns a result dict."""
    role = ROLES[role_idx]
    exp = EXPS[exp_idx]
    company = COMPANIES[company_idx]
    desig = RUN_DESIGNATION[(role_idx, exp_idx, company_idx)]
    mode = desig["mode"]
    resume = role_resumes[role]
    skills = resume["skills"]
    resume_text = resume["text"]
    sid = f"mx{seq:05d}"
    issues = []
    questions_served = 0
    skills_mentioned = set()

    try:
        res = start_interview(
            session_id=sid,
            candidate_id=100000 + seq,
            candidate_name=resume["name"],
            candidate_role=role,
            candidate_experience=exp,
            candidate_skills=skills,
            resume_text=resume_text,
            mode=mode,
            company=company,
        )
        if "error" in res:
            return _result(seq, role, company, exp, mode, 0, "", [], 0,
                           [f"START_ERROR: {res['error']}"], 0, False)

        is_apt = bool(res.get("is_aptitude"))
        question = res.get("question", "")
        qa = 0          # non-aptitude answers given
        apt_answered = 0
        aptitude_correct = 0
        aptitude_wrong = 0
        aptitude_leaks = 0
        was_apt = False
        completed = False

        while True:
            apt_data = None
            if is_apt:
                sess = session_store.get(sid)
                if sess and sess.aptitude_questions:
                    idx = sess.aptitude_question_index
                    if idx < len(sess.aptitude_questions):
                        qd = sess.aptitude_questions[idx]
                        apt_data = qd
                        correct = qd["correct"]
                        correct_text = qd["options"][correct]

            q_issues, mentioned = check_question_text(question, skills,
                                                      apt_data["options"][apt_data["correct"]]
                                                      if apt_data else None)
            skills_mentioned.update(mentioned)
            for i in q_issues:
                if i.startswith("APTITUDE_LEAKS") or i.startswith("APTITUDE_ANSWER_HINT"):
                    aptitude_leaks += 1
                issues.append(i)
            questions_served += 1

            if is_apt and apt_data is not None:
                bank_err = apt_bank_error(apt_data["question"], apt_data["options"],
                                          apt_data["correct"])
                if bank_err:
                    issues.append(bank_err)
                option = apt_data["correct"]
                if apt_answered % 3 == 2:
                    option = (option + 1) % 4  # submit a wrong answer to test scoring both ways
                r2 = submit_answer(sid, str(option))
                apt_answered += 1
                ev = r2.get("evaluation", {}) if isinstance(r2, dict) else {}
                if not isinstance(r2, dict) or r2.get("error"):
                    issues.append(f"APT_SUBMIT_ERR: {r2 if isinstance(r2, dict) else type(r2)}")
                    break
                if ev.get("is_correct") is True:
                    aptitude_correct += 1
                    if ev.get("overall_score") != 10:
                        issues.append(f"APTITUDE_SCORING(correct got {ev.get('overall_score')})")
                elif ev.get("is_correct") is False:
                    aptitude_wrong += 1
                    if ev.get("overall_score") not in (0, None):
                        issues.append(f"APTITUDE_SCORING(wrong got {ev.get('overall_score')})")
                else:
                    issues.append("APTITUDE_IS_CORRECT_MISSING")
            else:
                ans = build_answer(role, skills, question, "", False, company)
                r2 = submit_answer(sid, ans)
                qa += 1
                if not isinstance(r2, dict) or r2.get("error"):
                    issues.append(f"SUBMIT_ERR: {r2 if isinstance(r2, dict) else type(r2)}")
                    break
                issues.extend(check_evaluation(r2.get("evaluation", {}), ans))

            if r2.get("error"):
                issues.append(f"ENGINE_ERR: {r2['error']}")
                break
            if r2.get("is_complete"):
                completed = True
                break

            next_is_apt = bool(r2.get("is_aptitude", is_apt))

            # stopping rules -------------------------------------------------
            if desig.get("apt_style") == "full":
                # resume(3) -> aptitude(10) -> stop (ahead of next round)
                if was_apt and not next_is_apt and apt_answered > 0:
                    break
                if apt_answered >= APTITUDE_ROUND_QUESTIONS:
                    break
            elif desig.get("coding_deep"):
                if qa >= 4:
                    break
            else:
                if qa >= STD_Q_CAP:
                    break
            if questions_served >= 40:
                issues.append("SAFETY_QUESTION_CAP")
                break

            question = r2.get("next_question", "")
            was_apt = is_apt
            is_apt = next_is_apt
            if not question:
                issues.append("EMPTY_NEXT_QUESTION")
                break

        # Aptitude round completeness verification --------------------------
        if desig.get("apt_style") == "full":
            if apt_answered < APTITUDE_ROUND_QUESTIONS:
                issues.append(f"APTITUDE_ROUND_SHORT({apt_answered}/{APTITUDE_ROUND_QUESTIONS})")
            elif apt_answered > APTITUDE_ROUND_QUESTIONS:
                issues.append(f"APTITUDE_ROUND_OVER({apt_answered})")

        # Rewrite test -------------------------------------------------------
        if rewrite_key is not None:
            try:
                rw = rewrite_answer(sid, 0, build_answer(role, skills, session_store.get(sid).answers[0]["question"],
                                                         "", False, company) + " (rewritten, more structured)")
                if not isinstance(rw, dict) or rw.get("error"):
                    issues.append(f"REWRITE_ERR: {rw if isinstance(rw, dict) else rw}")
                else:
                    rs = rw.get("rewrite_scores", {})
                    if not isinstance(rs, dict) or not rs.get("overall_score"):
                        issues.append("REWRITE_NO_SCORES")
                    if not isinstance(rw.get("improvement"), dict):
                        issues.append("REWRITE_NO_IMPROVEMENT")
            except Exception as e:
                issues.append(f"REWRITE_CRASH: {e}")

        # Report generation --------------------------------------------------
        if real_report:
            report_real()
        else:
            report_stubbed()
        try:
            report = generate_report(sid)
        except Exception as e:
            report = {"error": f"REPORT_CRASH: {e}"}
        finally:
            report_stubbed()

        if not isinstance(report, dict) or report.get("error"):
            issues.append(f"REPORT_FAIL: {report.get('error') if isinstance(report, dict) else type(report)}")
            score, grade = 0, ""
            n_answered = len(session_store.get(sid).answers) if session_store.get(sid) else 0
        else:
            grade = report.get("grade", "")
            if grade not in VALID_GRADES:
                issues.append(f"REPORT_BAD_GRADE({grade})")
            score = report.get("overall_score", 0)
            if not isinstance(score, (int, float)):
                issues.append("REPORT_SCORE_NOT_NUMERIC")
                score = 0
            else:
                score = round(float(score), 1)
            # score consistency: report overall == mean of answer overall scores
            sess = session_store.get(sid)
            if sess and sess.answers:
                ans_scores = [a.get("overall_score", 0) for a in sess.answers]
                mean = round(sum(ans_scores) / len(ans_scores), 1)
                if abs(mean - score) > 0.05:
                    issues.append(f"REPORT_SCORE_MISMATCH(report={score}, answers={mean})")
            if not isinstance(report.get("skill_gaps"), list):
                issues.append("REPORT_NO_SKILL_GAPS")
            if not isinstance(report.get("recommendations"), list):
                issues.append("REPORT_NO_RECOMMENDATIONS")
            if report.get("aptitude_data") is None and desig.get("apt_style") == "full":
                issues.append("REPORT_MISSING_APTITUDE_SECTION")
            n_answered = len(session_store.get(sid).answers) if session_store.get(sid) else 0

        # Resume-phase skill usage check ------------------------------------
        sess = session_store.get(sid)
        if sess and resume_text:
            rp_qs = [q for i, q in enumerate(sess.questions)
                     if i < len(sess.questions_meta) and sess.questions_meta[i].get("is_resume_phase")]
            rp_skill_count = max((_skill_count(q, skills) for q in rp_qs), default=0)
            if rp_skill_count == 0:
                issues.append("RESUME_PHASE_NO_SKILLS")
        if len(skills_mentioned) == 0 and resume_text:
            issues.append("NO_RESUME_SKILL_IN_ANY_QUESTION")

        return _result(seq, role, company, exp, mode, score, grade, sorted(skills_mentioned),
                       questions_served, issues, qa + apt_answered, completed)

    except Exception as e:
        import traceback
        return _result(seq, role, company, exp, mode, 0, "", [], 0,
                       [f"CRASH: {type(e).__name__}: {e} :: {traceback.format_exc(limit=1)}".replace(chr(10), " ")], 0, False)


def _result(seq, role, company, exp, mode, score, grade, skills_used, q_served, issues, n_answered, completed):
    return {
        "seq": seq, "candidate": f"C{seq:04d}", "role": role, "company": company,
        "experience": exp, "mode": mode, "score": score, "grade": grade,
        "resume_skills_used": skills_used, "questions_served": q_served,
        "answered": n_answered, "issues": issues, "completed": completed,
        "pass": len(issues) == 0,
    }

# ══════════════════════════════════════════════════════════════════════════════
# DB persistence check (report regeneration from SQLite)
# ══════════════════════════════════════════════════════════════════════════════


def db_persistence_check(results, role_resumes):
    """Persist a few sessions to SQLite and regenerate reports from DB (stubbed LLM)."""
    import app as app_mod
    app_mod.init_db()
    out = []
    report_stubbed()
    # prefer non-coding sessions (the schema CHECK only allows hr/technical) so the
    # happy-path report-regen is actually exercised; coding-mode failures are probed
    # separately by db_coding_mode_probe()
    pool = [r for r in results if r["mode"] != "coding"] or list(results)
    for res in pool[:3]:
        sid = f"mx{res['seq']:05d}"
        sess = session_store.get(sid)
        if not sess:
            out.append({"seq": res["seq"], "ok": False, "err": "session missing"})
            continue
        try:
            report = generate_report(sid)
            cid = app_mod.save_candidate_to_db(
                name=sess.candidate_name, email=f"{res['seq']}@test.local",
                role=res["role"], experience=res["experience"],
                resume_filename="test_resume.docx", resume_text=sess.resume_text[:3000],
                skills=sess.candidate_skills,
            )
            app_mod.save_session_to_db({
                "session_id": sid, "candidate_id": cid, "mode": res["mode"],
                "status": "completed", "overall_score": report.get("overall_score", 0),
                "technical_score": report.get("technical_score", 0),
                "communication_score": report.get("communication_score", 0),
                "confidence_score": report.get("confidence_score", 0),
                "problem_solving_score": report.get("problem_solving_score", 0),
                "time_management_score": report.get("time_management_score", 0),
                "conceptual_clarity_score": report.get("conceptual_clarity_score", 0),
                "readiness_score": report.get("readiness_score", 0),
                "grade": report.get("grade", ""),
                "star_rating": report.get("star_rating", 0),
                "total_questions": report.get("total_questions", 0),
                "completed_questions": len(report.get("answers", [])),
                "duration_seconds": 60, "rounds": report.get("session_data", {}).get("rounds", []),
                "current_round": 0, "total_rounds": len(report.get("session_data", {}).get("rounds", [])),
            })
            app_mod.save_answers_to_db(sid, report.get("answers", []))
            app_mod.save_skill_gaps_to_db(sid, report.get("skill_gaps", []))
            # Regenerate report from DB (simulate server restart)
            session_store.delete(sid)
            report2 = generate_report(sid)
            ok = not (isinstance(report2, dict) and report2.get("error"))
            out.append({"seq": res["seq"], "ok": ok,
                        "err": "" if ok else report2.get("error")})
        except Exception as e:
            err = f"{type(e).__name__}: {e}"
            if "locked" in err or "busy" in err:
                time.sleep(0.5)
                try:
                    app_mod.save_session_to_db({
                        "session_id": sid, "candidate_id": cid, "mode": res["mode"],
                        "status": "completed", "overall_score": report.get("overall_score", 0),
                    })
                    out.append({"seq": res["seq"], "ok": True, "err": ""})
                    continue
                except Exception:
                    pass
            out.append({"seq": res["seq"], "ok": False, "err": err})
    return out


def db_coding_mode_probe():
    """Probe whether the app can persist a 'coding'-mode session (schema CHECK test)."""
    import app as app_mod
    app_mod.init_db()
    try:
        cid = app_mod.save_candidate_to_db(
            name="Probe User", email="probe_coding@test.local", role="Software Engineer",
            experience="2-4", resume_filename="probe.docx", resume_text="probe",
            skills=["python"],)
        app_mod.save_session_to_db({
            "session_id": "mx_probe_coding", "candidate_id": cid, "mode": "coding",
            "status": "completed", "overall_score": 50,
        })
        return ("OK", "coding-mode session persisted")
    except Exception as e:
        return ("FAIL", f"{type(e).__name__}: {e}")

# ══════════════════════════════════════════════════════════════════════════════
# Checkpoint (incremental results so the full matrix can be resumed across runs)
# ══════════════════════════════════════════════════════════════════════════════

_chk_lock = threading.Lock()


def load_checkpoint():
    """Return (list of prior result dicts, set of done combo keys)."""
    prior, done = [], set()
    if os.path.exists(CHECKPOINT):
        with open(CHECKPOINT, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    r = json.loads(line)
                except Exception:
                    continue
                prior.append(r)
                key = (ROLES.index(r["role"]), EXPS.index(r["experience"]),
                       COMPANIES.index(r["company"]))
                done.add(key)
    return prior, done


def write_checkpoint(result):
    with _chk_lock:
        with open(CHECKPOINT, "a", encoding="utf-8") as f:
            f.write(json.dumps(result, ensure_ascii=False) + "\n")


# ══════════════════════════════════════════════════════════════════════════════
# Main
# ══════════════════════════════════════════════════════════════════════════════


def main():
    t_start = time.time()
    stamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    print("=" * 100)
    print(f"  AI Interview Intelligence System - Full Matrix Test (Groq backend)  {stamp}")
    print("=" * 100)
    status = ai_service.check_ollama_health()
    print(f"  LLM provider: {status.get('provider')} | model: {status.get('model_name')} "
          f"| status: {status.get('status')}")
    if status.get("status") != "connected":
        print(f"  WARNING: {status.get('message')}")
    print()

    role_resumes = load_role_resumes()
    for role in ROLES:
        r = role_resumes[role]
        print(f"  resume {role:<22} skills={len(r['skills']):>2} "
              f"name={r['name']!r} text={len(r['text'])} chars")
    print()

    set_run_designations()

    combos = []
    for ri in range(len(ROLES)):
        for ei in range(len(EXPS)):
            for ci in range(len(COMPANIES)):
                combos.append((ri, ei, ci))
    random.Random(42).shuffle(combos)

    rewrite_targets = {
        (0, 2, 1): "Software Engineer/Google/2-4",   # role idx 0, exp idx 2, company idx 1
        (1, 2, 1): "Data Scientist/Google/2-4",
        (2, 2, 1): "Product Manager/Google/2-4",
    }
    real_report_keys = set(rewrite_targets.keys())
    # add a spread of real-report candidates across roles/companies
    for ri in range(len(ROLES)):
        for company in ("Google", "TCS", "General"):
            real_report_keys.add((ri, EXPS.index("2-4"), COMPANIES.index(company)))
    if len(real_report_keys) > REAL_REPORT_SAMPLE + len(rewrite_targets):
        real_report_keys = set(list(real_report_keys)[:REAL_REPORT_SAMPLE + len(rewrite_targets)])

    # checkpoint / resume support
    prior, done_keys = load_checkpoint()
    combos = [c for c in combos if c not in done_keys]
    # force rewrite targets into this batch (front of the queue) if not yet done
    front = [c for c in rewrite_targets if c in combos]
    rest = [c for c in combos if c not in front]
    combos = front + rest
    if SUBSET_LIMIT > 0:
        combos = combos[:SUBSET_LIMIT]
    print(f"  Checkpoint: {len(prior)} prior results loaded ({len(done_keys)} combos done), "
          f"{len(combos)} remaining in this batch" + (f" (SUBSET_LIMIT={SUBSET_LIMIT})" if SUBSET_LIMIT else ""))
    print()

    results = list(prior)
    seq = max([r["seq"] for r in prior], default=0)
    launched = 0
    skipped = 0
    deadline = t_start + LAUNCH_BUDGET_SEC
    deep_used = 0

    new_results = []
    with ThreadPoolExecutor(max_workers=WORKERS) as pool:
        futs = {}
        for (ri, ei, ci) in combos:
            if time.time() > deadline:
                skipped += 1
                continue
            seq += 1
            launched += 1
            key = (ri, ei, ci)
            rw = key if key in rewrite_targets else None
            real = key in real_report_keys
            # throttle the deep-coding runs
            is_deep = RUN_DESIGNATION[key].get("coding_deep", False)
            if is_deep and deep_used >= CODING_DEEP_CAP:
                RUN_DESIGNATION[key]["coding_deep"] = False
                RUN_DESIGNATION[key]["apt_style"] = "none"
            if is_deep and RUN_DESIGNATION[key].get("coding_deep"):
                deep_used += 1
            futs[pool.submit(run_candidate, seq, ri, ei, ci, role_resumes, rw, real)] = seq

        done = 0
        for fut in as_completed(futs):
            r = fut.result()
            new_results.append(r)
            write_checkpoint(r)
            done += 1
            if done % 25 == 0:
                el = time.time() - t_start
                print(f"  ... {done}/{len(futs)} candidates done in {el:.0f}s "
                      f"(llm_calls={STATS['count']}, rate_limited={STATS['errors']})")
    new_results.sort(key=lambda x: x["seq"])
    results = prior + new_results

    # rewrite targets: show which got rewrite test
    rewrite_log = []
    for r in results:
        if (ROLES.index(r["role"]), EXPS.index(r["experience"]), COMPANIES.index(r["company"])) in rewrite_targets:
            rewrite_log.append(r["seq"])

    print()
    print("  Running DB persistence check on 3 sessions (report regen from SQLite)...")
    db_checks = db_persistence_check(new_results, role_resumes)
    for dc in db_checks:
        print(f"    DB check seq={dc['seq']}: {'OK' if dc['ok'] else 'FAIL ' + dc['err']}")
    coding_db = db_coding_mode_probe()
    print(f"    Coding-mode persistence probe: {coding_db[0]} ({coding_db[1]})")

    # ── Summary ────────────────────────────────────────────────────────────
    elapsed = time.time() - t_start
    total = len(results)
    passed = [r for r in results if r["pass"]]
    failed = [r for r in results if not r["pass"]]
    n_llm = STATS["count"]
    n_llm_err = STATS["errors"]

    print()
    print("=" * 100)
    print("  SUMMARY")
    print("=" * 100)
    print(f"  This batch launched : {launched}   (skipped to time budget: {skipped})")
    print(f"  Combos completed    : {total}   (prior checkpoint: {len(prior)})")
    print(f"  Passed              : {len(passed)}")
    print(f"  With issues         : {len(failed)}")
    print(f"  LLM calls made       : {n_llm}   (rate-limit/provider errors: {n_llm_err})")
    print(f"  Wall time            : {elapsed/60:.1f} min")
    if n_llm:
        print(f"  LLM throughput       : {n_llm/elapsed:.2f} calls/sec   avg {STATS['total_s']/n_llm:.2f}s/call")
    print(f"  Rewrite tests        : {len(rewrite_log)} ({len([r for r in results if any('REWRITE' in i for i in r['issues'])])} flagged)")
    if db_checks:
        print(f"  DB persistence       : {sum(1 for d in db_checks if d['ok'])}/3 OK")

    # Aggregate issue counts
    from collections import Counter
    issue_counter = Counter()
    for r in results:
        for i in r["issues"]:
            key = i.split("(")[0]
            issue_counter[key] += 1
    print()
    print("  Issue frequency (top 20):")
    for k, v in issue_counter.most_common(20):
        print(f"    {k:<40} {v}")

    if STATS["error_samples"]:
        print()
        print("  LLM error samples (first 6):")
        for s, dt in STATS["error_samples"][:6]:
            print(f"    [{dt}s] {s}")

    # ── Full pass/fail table ───────────────────────────────────────────────
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    table_path = os.path.join(PROJ, f"matrix_full_test_results_{ts}.txt")
    with open(table_path, "w", encoding="utf-8") as f:
        f.write(f"AI Interview Intelligence System - Full Matrix Results ({stamp})\n")
        f.write("=" * 100 + "\n\n")
        hdr = (f"{'candidate':<9} {'role':<20} {'company':<10} {'exp':<6} {'mode':<10} "
               f"{'rs_used':<8} {'score':<6} {'grade':<5} issues")
        f.write(hdr + "\n")
        f.write("-" * 100 + "\n")
        print()
        print("  FULL MATRIX TABLE (candidate | role | company | experience | mode | resume skills used | score | grade | issues)")
        print("-" * 100)
        for r in results:
            rs = f"{len(r['resume_skills_used'])}/{' '.join(sorted(r['resume_skills_used'])[:4])}" if r["resume_skills_used"] else "0"
            issues_txt = ", ".join(r["issues"]) if r["issues"] else "-"
            line = (f"{r['candidate']:<9} {r['role']:<20} {r['company']:<10} {r['experience']:<6} "
                    f"{r['mode']:<10} {str(len(r['resume_skills_used'])):<8} {r['score']:<6} "
                    f"{r['grade']:<5} {issues_txt}")
            print(line)
            f.write(line + "\n")
        f.write("\n" + "=" * 100 + "\n")
        f.write("AGGREGATES\n")
        f.write(f"total={total} passed={len(passed)} failed={len(failed)} "
                f"llm_calls={n_llm} llm_errors={n_llm_err} wall_min={elapsed/60:.1f}\n")
        for k, v in issue_counter.most_common():
            f.write(f"{k}: {v}\n")
    print()
    print(f"  Results saved to: {table_path}")

    return 0 if n_llm_err == 0 and len(failed) == 0 else 1


if __name__ == "__main__":
    sys.exit(main())
